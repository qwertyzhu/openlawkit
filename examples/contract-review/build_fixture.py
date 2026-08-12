#!/usr/bin/env python3
"""Create the entirely fictional DOCX used by the public example and tests."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


CLAUSES = [
    "本材料完全虚构，仅用于 OpenLawKit 软件测试，不构成真实交易文件或法律意见。",
    "甲方：青岚文化科技（杭州）有限公司（虚构）",
    "乙方：远川数智服务（宁波）有限公司（虚构）",
    "第一条 乙方为甲方提供虚构的内容管理系统设计、开发和测试服务。",
    "第二条 服务期限自2026年3月1日至2026年5月31日。",
    "第三条 服务费总额为人民币壹拾万元整（¥100,000.00）。乙方开具发票后，甲方支付人民币壹拾贰万元整（¥120,000.00）。",
    "第四条 乙方提交成果后，甲方应在两个自然日内验收；甲方未在两个自然日内提出异议的，视为验收合格。验收标准以甲方满意为准。",
    "第五条 项目成果及相关知识产权均归乙方所有，甲方仅可在本项目内部使用，合同终止后应立即停止使用。",
    "第六条 双方对履约中获悉的商业秘密承担保密义务，保密义务仅在本合同有效期内存续。",
    "第七条 乙方因任何违约承担的全部责任，以甲方最近一期已支付服务费的百分之十为上限。",
    "第八条 任何一方提前七日书面通知即可解除本合同，已完成部分按双方另行协商的金额结算。",
    "第九条 因本合同产生的争议，双方应提交当地仲裁委员会仲裁。",
    "第十条 本合同一式两份，双方各执一份，自盖章之日起生效。",
    "（以下无正文，为完全虚构测试签署页）",
    "甲方（盖章）：________________    乙方（盖章）：________________",
]


def build(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("信息技术服务合同（完全虚构测试材料）")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "SimHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for clause in CLAUSES:
        paragraph = document.add_paragraph(clause)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.35

    properties = document.core_properties
    properties.title = "OpenLawKit 完全虚构合同审查测试材料"
    properties.subject = "Public fictional fixture"
    properties.author = "OpenLawKit Contributors"
    properties.last_modified_by = "OpenLawKit Contributors"
    properties.comments = "All names, facts and clauses are fictional."
    fixed = datetime(2026, 1, 1, tzinfo=timezone.utc)
    properties.created = fixed
    properties.modified = fixed
    properties.revision = 1

    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)
    print(f"created: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
